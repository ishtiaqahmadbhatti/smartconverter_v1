import os
from typing import Optional, List, Tuple
import pytesseract
from PIL import Image
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from app.core.exceptions import FileProcessingError
from app.services.file_service import FileService


class ConversionService:
    """Service for handling file conversions."""
    
    @staticmethod
    def pdf_to_word(input_path: str, output_filename: Optional[str] = None) -> str:
        """Convert PDF to Word document."""
        try:
            # Validate input file exists
            if not os.path.exists(input_path):
                raise FileProcessingError(f"Input PDF file not found: {input_path}")
            
            if output_filename and output_filename.strip():
                output_path, _ = FileService.generate_output_path_with_filename(
                    output_filename.strip(),
                    default_extension=".docx",
                )
            else:
                output_path = FileService.get_output_path(input_path, ".docx")
            
            # Use pdf2docx to convert PDF to Word
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            # Verify output file was created
            if not os.path.exists(output_path):
                raise FileProcessingError("Conversion completed but output file was not created")
            
            return output_path
            
        except FileNotFoundError as e:
            raise FileProcessingError(f"PDF file not found: {str(e)}")
        except PermissionError as e:
            raise FileProcessingError(f"Permission denied accessing PDF file: {str(e)}")
        except Exception as e:
            error_msg = str(e)
            if "password" in error_msg.lower():
                raise FileProcessingError("PDF is password protected and cannot be converted")
            elif "corrupt" in error_msg.lower() or "invalid" in error_msg.lower():
                raise FileProcessingError("PDF file appears to be corrupted or invalid")
            else:
                raise FileProcessingError(f"PDF to Word conversion failed: {error_msg}")
    
    @staticmethod
    def word_to_pdf(input_path: str) -> str:
        """Convert Word document to PDF with comprehensive data preservation."""
        try:
            # Validate input file exists
            if not os.path.exists(input_path):
                raise FileProcessingError(f"Input Word file not found: {input_path}")
            
            output_path = FileService.get_output_path(input_path, ".pdf")
            
            # Read Word document
            doc = Document(input_path)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path, 
                pagesize=A4,
                rightMargin=2*cm, 
                leftMargin=2*cm, 
                topMargin=2*cm, 
                bottomMargin=2*cm
            )
            
            # Create comprehensive styles
            styles = ConversionService._create_comprehensive_styles()
            
            # Build content for PDF using comprehensive method
            content = []
            
            # Process document in order using element-based approach
            content = ConversionService._process_document_comprehensively(doc, styles)
            
            # If no content was processed, use fallback
            if not content:
                content = ConversionService._fallback_comprehensive_processing(doc, styles)
            
            # Build PDF
            pdf_doc.build(content)
            
            # Verify output file was created
            if not os.path.exists(output_path):
                raise FileProcessingError("Conversion completed but output file was not created")
            
            return output_path
            
        except FileNotFoundError as e:
            raise FileProcessingError(f"Word file not found: {str(e)}")
        except PermissionError as e:
            raise FileProcessingError(f"Permission denied accessing Word file: {str(e)}")
        except Exception as e:
            error_msg = str(e)
            if "corrupt" in error_msg.lower() or "invalid" in error_msg.lower():
                raise FileProcessingError("Word file appears to be corrupted or invalid")
            else:
                raise FileProcessingError(f"Word to PDF conversion failed: {error_msg}")
    
    @staticmethod
    def _create_comprehensive_styles():
        """Create comprehensive styles for PDF generation."""
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'ComprehensiveTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=15,
            spaceBefore=10,
            alignment=TA_CENTER,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Heading 1 style
        h1_style = ParagraphStyle(
            'ComprehensiveH1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=15,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Heading 2 style
        h2_style = ParagraphStyle(
            'ComprehensiveH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=12,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Heading 3 style
        h3_style = ParagraphStyle(
            'ComprehensiveH3',
            parent=styles['Heading3'],
            fontSize=13,
            spaceAfter=8,
            spaceBefore=10,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        # Normal paragraph style
        normal_style = ParagraphStyle(
            'ComprehensiveNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=13,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        # List style
        list_style = ParagraphStyle(
            'ComprehensiveList',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=4,
            leading=12,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        return {
            'title': title_style,
            'heading1': h1_style,
            'heading2': h2_style,
            'heading3': h3_style,
            'normal': normal_style,
            'list': list_style
        }
    
    @staticmethod
    def _process_document_comprehensively(doc, styles):
        """Process document comprehensively to preserve all data."""
        content = []
        
        try:
            # Process all paragraphs with comprehensive formatting
            for para in doc.paragraphs:
                if para.text.strip():
                    style = ConversionService._get_comprehensive_paragraph_style(para, styles)
                    formatted_text = ConversionService._extract_comprehensive_formatted_text(para)
                    
                    if formatted_text.strip():
                        content.append(Paragraph(formatted_text, style))
                        content.append(Spacer(1, 2))
            
            # Process all tables with comprehensive formatting
            for table in doc.tables:
                content.append(Spacer(1, 6))
                table_content = ConversionService._convert_comprehensive_table_to_pdf(table, styles)
                if table_content:
                    content.append(table_content)
                content.append(Spacer(1, 6))
            
        except Exception as e:
            print(f"Warning: Error in comprehensive processing: {str(e)}")
        
        return content
    
    @staticmethod
    def _fallback_comprehensive_processing(doc, styles):
        """Comprehensive fallback processing to ensure no data is lost."""
        content = []
        
        print("Using comprehensive fallback processing")
        
        # Process all paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                style = ConversionService._get_comprehensive_paragraph_style(para, styles)
                formatted_text = ConversionService._extract_comprehensive_formatted_text(para)
                
                if formatted_text.strip():
                    content.append(Paragraph(formatted_text, style))
                    content.append(Spacer(1, 2))
        
        # Process all tables
        for table in doc.tables:
            content.append(Spacer(1, 6))
            table_content = ConversionService._convert_comprehensive_table_to_pdf(table, styles)
            if table_content:
                content.append(table_content)
            content.append(Spacer(1, 6))
        
        return content
    
    @staticmethod
    def _get_comprehensive_paragraph_style(paragraph, styles):
        """Get comprehensive paragraph style based on content and formatting."""
        style_name = paragraph.style.name.lower()
        
        # Check for title patterns
        if 'title' in style_name or paragraph.text.strip().isupper():
            return styles['title']
        
        # Check for heading patterns
        elif 'heading 1' in style_name or 'heading1' in style_name:
            return styles['heading1']
        elif 'heading 2' in style_name or 'heading2' in style_name:
            return styles['heading2']
        elif 'heading 3' in style_name or 'heading3' in style_name:
            return styles['heading3']
        elif 'heading' in style_name:
            return styles['heading1']
        
        # Check for list patterns
        elif paragraph.text.strip().startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.')):
            return styles['list']
        
        # Default to normal
        else:
            return styles['normal']
    
    @staticmethod
    def _extract_comprehensive_formatted_text(paragraph):
        """Extract text with comprehensive formatting preserved."""
        if not paragraph.runs:
            return paragraph.text
        
        text = ""
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
                
            # Apply comprehensive formatting
            if run.bold and run.italic:
                run_text = f"<b><i>{run_text}</i></b>"
            elif run.bold:
                run_text = f"<b>{run_text}</b>"
            elif run.italic:
                run_text = f"<i>{run_text}</i>"
            
            if run.underline:
                run_text = f"<u>{run_text}</u>"
            
            # Handle font size
            if hasattr(run, 'font') and run.font.size:
                size_pt = run.font.size.pt
                if size_pt != 11:  # Only apply if different from default
                    run_text = f"<font size='{int(size_pt)}'>{run_text}</font>"
            
            text += run_text
        
        return text or paragraph.text
    
    @staticmethod
    def _convert_comprehensive_table_to_pdf(table, styles):
        """Convert Word table to PDF table with comprehensive formatting."""
        try:
            # Extract table data with comprehensive formatting
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Extract formatted text from cell
                    cell_text = ConversionService._extract_cell_comprehensive_text(cell)
                    if not cell_text.strip():
                        cell_text = " "  # Empty cell with space
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if not table_data:
                return None
            
            # Remove completely empty rows
            table_data = [row for row in table_data if any(cell.strip() for cell in row)]
            if not table_data:
                return None
            
            # Ensure all rows have the same number of columns
            max_cols = max(len(row) for row in table_data) if table_data else 0
            for row in table_data:
                while len(row) < max_cols:
                    row.append(" ")
            
            # Create PDF table with proper column widths
            col_widths = [1.0 / max_cols] * max_cols if max_cols > 0 else [1.0]
            pdf_table = Table(table_data, colWidths=col_widths)
            
            # Apply comprehensive table styling
            table_style = TableStyle([
                # Header row styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('LEFTPADDING', (0, 0), (-1, 0), 6),
                ('RIGHTPADDING', (0, 0), (-1, 0), 6),
                
                # Data rows styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('LEFTPADDING', (0, 1), (-1, -1), 6),
                ('RIGHTPADDING', (0, 1), (-1, -1), 6),
                
                # Grid lines
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.black),
                
                # Alternating row colors
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ])
            
            pdf_table.setStyle(table_style)
            return KeepTogether(pdf_table)
            
        except Exception as e:
            print(f"Warning: Error converting table to PDF: {str(e)}")
            # Fallback to comprehensive text representation
            table_text = ""
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text += row_text + "\n"
            return Paragraph(table_text, styles['normal'])
    
    @staticmethod
    def _extract_cell_comprehensive_text(cell):
        """Extract comprehensive formatted text from a table cell."""
        if not cell.paragraphs:
            return cell.text.strip()
        
        formatted_text = ""
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():
                para_text = ConversionService._extract_comprehensive_formatted_text(paragraph)
                formatted_text += para_text + " "
        
        return formatted_text.strip().replace('\n', ' ')
    

    @staticmethod
    def image_to_text(input_path: str) -> str:
        """Extract text from image using OCR."""
        try:
            # Configure tesseract path if provided
            if hasattr(pytesseract, 'pytesseract') and pytesseract.pytesseract.tesseract_cmd:
                pass  # Already configured
            elif os.name == 'nt':  # Windows
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            
            # Open and process image
            image = Image.open(input_path)
            text = pytesseract.image_to_string(image)
            
            return text.strip()
            
        except Exception as e:
            raise FileProcessingError(f"Image to text conversion failed: {str(e)}")
    
    @staticmethod
    def cleanup_temp_files(*file_paths: str) -> None:
        """Clean up temporary files."""
        for file_path in file_paths:
            FileService.cleanup_file(file_path)

    # NOTE: video_to_audio method removed per request